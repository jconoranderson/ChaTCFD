from __future__ import annotations

import io
from pathlib import Path
from typing import List, Optional

import docx
import fitz  # PyMuPDF
from docx.oxml.ns import qn

from .providers import ModelProvider
from .rag import RAGStore
from .settings import Settings


class BIPService:
    def __init__(self, settings: Settings, provider: ModelProvider, rag_store: RAGStore):
        self.settings = settings
        self.provider = provider
        self.rag_store = rag_store
        self._few_shot_examples = self._load_examples()

    def _examples_dir(self) -> Path:
        return Path(self.settings.bip_examples_dir)

    def _load_examples(self) -> List[str]:
        directory = self._examples_dir()
        if not directory.exists():
            return []
        examples: List[str] = []
        for path in sorted(directory.glob("*.txt")):
            content = path.read_text().strip()
            if content:
                examples.append(content)
        return examples

    def build_prompt(
        self,
        name: str,
        age: int,
        diagnosis: str,
        behavior: str,
        setting: str,
        trigger: str,
        notes: Optional[str],
        fba_text: Optional[str],
    ) -> str:
        examples_block = "\n---\n".join(self._few_shot_examples[:3]) if self._few_shot_examples else ""

        user_profile = (
            "Student Profile:\n"
            f"- Name: {name}\n"
            f"- Age: {age}\n"
            f"- Diagnosis: {diagnosis}\n"
            f"- Behavior: {behavior}\n"
            f"- Setting: {setting}\n"
            f"- Trigger: {trigger}\n"
        )
        if notes:
            user_profile += f"- Notes: {notes}\n"

        if fba_text:
            user_profile += f"\nFunctional Behavior Assessment Summary:\n{fba_text.strip()}\n"

        policy_nodes = []
        try:
            policy_nodes = self.rag_store.retrieve("bip_policies", user_profile)
        except Exception:
            policy_nodes = []

        policy_context = "\n\n".join(
            node.node.get_content().strip()
            for node in policy_nodes
            if node.node and node.node.get_content().strip()
        )

        prompt = (
            "You are a certified behavior analyst creating a Behavior Intervention Plan (BIP) "
            "for The Center for Discovery. Use people-first, observable, measurable language. "
            "Ensure replacement behaviors are functionally equivalent to the target behavior."
        )

        prompt += (
            "\n\nFollow these guidelines when writing the plan:"
            "\n- Adhere to New York State OPWDD and The Center for Discovery standards."
            "\n- Avoid mentalistic explanations (e.g., 'the student wants attention')."
            "\n- Provide goals that are measurable with clear criteria."
            "\n- Include safety precautions when relevant."
        )

        if examples_block:
            prompt += f"\n\n[REFERENCE EXAMPLES]\n{examples_block}"

        if policy_context:
            prompt += f"\n\n[POLICY CONTEXT]\n{policy_context}"

        prompt += (
            "\n\n[NEW REQUEST]\n"
            f"{user_profile}\n"
            "\nPlease produce a complete BIP that includes:\n"
            "- FBA Summary\n"
            "- Operational Definition\n"
            "- Replacement Behaviors\n"
            "- Prevention Strategies\n"
            "- Reinforcement Plan\n"
            "- Data Collection Method\n"
            "- Crisis/Safety Plan if applicable\n"
            "- Three short-term goals and one long-term goal with measurable criteria\n"
        )
        return prompt

    def generate(self, prompt: str, model_override: Optional[str] = None) -> str:
        return self.provider.chat(
            messages=[{"role": "user", "content": prompt}],
            model=model_override,
        )

    @staticmethod
    def extract_text_from_upload(filename: str, content: bytes) -> Optional[str]:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return BIPService._extract_pdf(content)
        if suffix == ".docx":
            return BIPService._extract_docx(content)
        if suffix == ".txt":
            return content.decode("utf-8", errors="ignore")
        return None

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        with fitz.open(stream=content, filetype="pdf") as doc:
            pages = [page.get_text() for page in doc]
        return "\n".join(pages)

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        file_like = io.BytesIO(content)
        document = docx.Document(file_like)
        lines: List[str] = []

        def extract_paragraph(paragraph: docx.text.paragraph.Paragraph) -> str:
            if paragraph.runs:
                bits = []
                for run in paragraph.runs:
                    text = run.text
                    if text:
                        bits.append(text)
                combined = "".join(bits)
                if combined:
                    return combined
            return paragraph.text

        for paragraph in document.paragraphs:
            text = extract_paragraph(paragraph).strip()
            if text:
                lines.append(text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_lines = []
                    for paragraph in cell.paragraphs:
                        text = extract_paragraph(paragraph).strip()
                        if text:
                            cell_lines.append(text)
                    text = " ".join(cell_lines)
                    if text:
                        lines.append(text)

        return "\n".join(lines)
